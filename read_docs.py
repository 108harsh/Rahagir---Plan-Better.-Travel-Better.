import docx
import os
import sys

# Force UTF-8 for stdout
sys.stdout.reconfigure(encoding='utf-8')

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Read Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Read Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading {file_path}: {e}"

files = [
    r"c:\Users\Lenovo\Desktop\CAPSTONE\Rahagir\Rahagir todolist.docx",
    r"c:\Users\Lenovo\Desktop\CAPSTONE\Rahagir\RAHAGIR_pitch_text.docx"
]

for f in files:
    print(f"\n--- Reading {os.path.basename(f)} ---")
    content = read_docx(f)
    print(content)
    print("------------------------------------------------\n")
